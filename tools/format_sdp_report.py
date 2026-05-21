from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


INPUT = Path(r"C:\Users\saiku\Downloads\SDP (1) (1).docx")
OUTPUT = Path(r"C:\Users\saiku\OneDrive\Desktop\Projects\SDP_Test\output\docx\SDP_report_format_checked.docx")


FRONT_MATTER = {
    "DECLARATION",
    "CERTIFICATE",
    "ABSTRACT",
    "ACKNOWLEDGEMENT",
    "CONTENTS",
    "LIST OF FIGURES",
    "LIST OF TABLES",
    "LIST OF ACRONYMS",
    "APPENDICES",
    "REFERENCES",
}


def set_run_font(run, size=12, bold=None, italic=None, underline=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline


def set_paragraph_format(paragraph, *, size=12, bold=None, italic=None, underline=None,
                         align=None, before=0, after=6, line_spacing=1.5):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    paragraph.style = paragraph.style
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, italic=italic, underline=underline)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_run_font(run, 12)


def is_chapter_marker(text):
    return bool(re.fullmatch(r"CHAPTER\s+\d+", text))


def is_section_heading(text):
    return bool(re.match(r"^\d+\.\d+\s+", text))


def is_subsection_heading(text):
    return bool(re.match(r"^\d+\.\d+\.\d+\s+", text))


def format_title_pages(paragraphs):
    for p in paragraphs[:32]:
        if not p.text.strip():
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0

    title_indexes = [i for i, p in enumerate(paragraphs[:32]) if p.text.strip().startswith("INTERPRETABLE AI SYSTEM")]
    for idx in title_indexes:
        set_paragraph_format(paragraphs[idx], size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8, line_spacing=1.0)

    for i, p in enumerate(paragraphs[:32]):
        text = p.text.strip()
        if text == "A project report on":
            set_paragraph_format(p, size=12, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=6, line_spacing=1.0)
        elif text.startswith("Submitted in partial fulfillment"):
            set_paragraph_format(p, size=14, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=3, line_spacing=1.0)
        elif text == "Bachelor’s Degree In CSE  by":
            set_paragraph_format(p, size=14, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=3, line_spacing=1.0)
        elif re.match(r"^[A-Z .]+\(22BCE\d+\)$", text):
            set_paragraph_format(p, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2, line_spacing=1.0)
        elif text == "SCHOOL OF COMPUTER SCIENCE AND ENGINEERING":
            set_paragraph_format(p, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, line_spacing=1.0)
        elif text == "JANUARY-MAY 2026":
            set_paragraph_format(p, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, line_spacing=1.0)

    date_seen = 0
    for p in paragraphs[:32]:
        if p.text.strip() == "JANUARY-MAY 2026":
            date_seen += 1
            if date_seen <= 2:
                p.add_run().add_break(WD_BREAK.PAGE)


def main():
    doc = Document(INPUT)

    for section in doc.sections:
        section.start_type = WD_SECTION.NEW_PAGE
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(3.81)
        section.right_margin = Cm(2.54)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.footer_distance = Cm(1.25)
        footer = section.footer
        footer.paragraphs[0].clear()
        add_page_number(footer.paragraphs[0])

    styles = doc.styles
    for style_name in ["Normal", "Body Text", "Normal (Web)"]:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            style.font.size = Pt(12)
            style.paragraph_format.line_spacing = 1.5
            style.paragraph_format.space_after = Pt(6)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        set_paragraph_format(paragraph, size=12, after=6, line_spacing=1.5)

        if text in FRONT_MATTER:
            size = 16 if text == "REFERENCES" else 14
            set_paragraph_format(
                paragraph,
                size=size,
                bold=True,
                underline=(text not in {"CONTENTS", "LIST OF FIGURES", "LIST OF TABLES", "LIST OF ACRONYMS"}),
                align=WD_ALIGN_PARAGRAPH.CENTER,
                before=12,
                after=12,
                line_spacing=1.5,
            )
        elif is_chapter_marker(text):
            set_paragraph_format(paragraph, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=6)
        elif text.startswith("CONCLUSION AND FUTURE WORK"):
            set_paragraph_format(paragraph, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=12)
        elif text in {"INTRODUCTION", "LITERATURE SURVEY", "SYSTEM OVERVIEW", "METHODOLOGY", "SYSTEM DESIGN", "IMPLEMENTATION", "RESULTS AND ANALYSIS", "RESULT AND ANALYSIS", "ADVANTAGES AND LIMITATIONS", "CODE"}:
            set_paragraph_format(paragraph, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
        elif is_subsection_heading(text):
            set_paragraph_format(paragraph, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=6)
        elif is_section_heading(text):
            set_paragraph_format(paragraph, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=10, after=6)

    format_title_pages(doc.paragraphs)

    for table in doc.tables:
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    set_paragraph_format(paragraph, size=12, after=0, line_spacing=1.0)

    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        p = doc.paragraphs[-1]._element
        p.getparent().remove(p)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
