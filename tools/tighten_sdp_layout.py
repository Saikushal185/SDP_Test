from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX = Path(r"C:\Users\saiku\OneDrive\Desktop\Projects\SDP_Test\output\docx\SDP_report_format_checked.docx")

BODY_MAJOR_HEADINGS = {
    "CHAPTER 1",
    "CHAPTER 2",
    "CHAPTER 3",
    "CHAPTER 4",
    "CHAPTER 5",
    "CHAPTER 6",
    "CHAPTER 7",
    "CHAPTER 8",
    "CHAPTER 9",
    "CHAPTER 10",
    "CHAPTER 11",
    "CONCLUSION AND FUTURE WORK",
    "APPENDICES",
    "REFERENCES",
}


def set_font(run, size=12, bold=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def clear_runs(paragraph):
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)


def format_toc_line(paragraph, text, toc_width_cm):
    clear_runs(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.15
    pf.tab_stops.clear_all()
    pf.tab_stops.add_tab_stop(Cm(toc_width_cm), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

    left, page = split_toc_entry(text)
    if is_subentry(left):
        pf.left_indent = Cm(1.35)
    else:
        pf.left_indent = Cm(0)

    run = paragraph.add_run(left)
    set_font(run, 14)
    if page:
        run = paragraph.add_run("\t" + page)
        set_font(run, 14)


def split_toc_entry(text):
    if "\t" in text:
        left, page = text.rsplit("\t", 1)
        return left.strip(), page.strip()
    return text.strip(), ""


def is_subentry(text):
    return bool(re.match(r"^\d+\.\d+", text.strip()))


def main():
    doc = Document(DOCX)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        pf = paragraph.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(3)
        pf.line_spacing = 1.15
        for run in paragraph.runs:
            if run.text:
                set_font(run, run.font.size.pt if run.font.size else 12)

    toc_start = None
    toc_end = None
    seen_references_entry = False
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() in {"CONTENTS", "TABLE OF CONTENTS"}:
            toc_start = i
        elif toc_start is not None and paragraph.text.strip().startswith("REFERENCES\t"):
            seen_references_entry = True
        elif toc_start is not None and seen_references_entry and not paragraph.text.strip():
            toc_end = i
            break

    if toc_start is not None:
        title = doc.paragraphs[toc_start]
        clear_runs(title)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(10)
        title.paragraph_format.line_spacing = 1.0
        run = title.add_run("TABLE OF CONTENTS")
        set_font(run, 16, bold=True)

        if toc_end is None:
            toc_end = min(toc_start + 70, len(doc.paragraphs))
        toc_width_cm = 14.65
        for paragraph in doc.paragraphs[toc_start + 1:toc_end]:
            text = paragraph.text.strip()
            if not text:
                paragraph.paragraph_format.space_after = Pt(0)
                continue
            format_toc_line(paragraph, text, toc_width_cm)

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if i <= 200:
            continue
        if text in BODY_MAJOR_HEADINGS or re.fullmatch(r"CHAPTER\s+\d+", text):
            paragraph.paragraph_format.page_break_before = True
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(6)

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
